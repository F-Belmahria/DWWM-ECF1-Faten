from pathlib import Path
import zipfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\faten\Documents\GitHub\DWWM-ECF1-Faten")
SOURCE_ODT = Path(r"C:\Users\faten\Downloads\1- Faten dossier_professionnel_version_DWWM .docx.odt")
OUT_DIR = ROOT / "output" / "documents"
IMG_DIR = ROOT / "tmp" / "photographie33_images"
OUT_DOCX = OUT_DIR / "Validation_travail_LA_Photographie33_Faten_Belmahria.docx"


BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY = "F2F4F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def style_run(run, bold=False, italic=False, color=None, size=None):
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)


def add_label_para(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(label)
    style_run(r, bold=True, color=DARK_BLUE)
    p.add_run(text)
    return p


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F7F7")
    set_cell_margins(cell, top=100, bottom=100, start=140, end=140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8.5)
    return table


def add_info_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    table.columns[0].width = Inches(1.9)
    table.columns[1].width = Inches(4.6)
    hdr = table.rows[0]
    hdr.cells[0].text = "Élément"
    hdr.cells[1].text = "Détail"
    for cell in hdr.cells:
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = DARK_BLUE
    for label, detail in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = detail
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
    return table


def add_three_col_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = DARK_BLUE
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[idx])
    return table


def add_image_with_caption(doc, image_path, caption, width=5.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    for r in cap.runs:
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GRAY


def extract_images():
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    mapping = {
        "media/image2.png": "wireframes.png",
        "media/image3.png": "figma_wireframes_detail.png",
        "media/image4.png": "maquette_accueil.png",
        "media/image5.png": "prototype_interactions.png",
        "media/image6.png": "palette_couleurs.png",
    }
    out = {}
    with zipfile.ZipFile(SOURCE_ODT) as z:
        for src, name in mapping.items():
            path = IMG_DIR / name
            path.write_bytes(z.read(src))
            out[src] = path
    return out


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    images = extract_images()
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("Validation du travail réalisé")
    style_run(r, bold=True, color=DARK_BLUE, size=22)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("Projet LA Photographie33 - site vitrine d'un photographe professionnel")
    style_run(r, italic=True, color=GRAY, size=12)

    add_info_table(doc, [
        ("Candidate", "BELMAHRIA Faten"),
        ("Titre visé", "Développeur Web et Web Mobile (DWWM)"),
        ("Activité-type", "Développer la partie front-end d'une application web"),
        ("Projet", "LA Photographie33 - site vitrine d'un photographe professionnel"),
        ("Organisation", "Conception Figma et développement réalisés en équipe"),
        ("Période attestée", "Juin 2026 - contributions GitHub du 5 au 9 juin 2026"),
        ("Technologies", "Figma, HTML5, SCSS, Bootstrap 5, Git et GitHub"),
    ])

    doc.add_heading("Fil conducteur", level=1)
    doc.add_paragraph(
        "J'ai participé avec l'équipe à la conception de l'expérience utilisateur et des écrans responsive. "
        "Pendant le développement, j'ai notamment pris en charge des éléments identifiables du site : "
        "l'illustration de contact et la page Blog."
    )

    doc.add_heading("Compétences mobilisées", level=2)
    add_three_col_table(doc, ["Compétence", "Travail réalisé", "Apport professionnel"], [
        ("UX/UI", "Organisation des contenus et maquettes responsive", "Parcours clair et cohérent"),
        ("Front-end", "HTML sémantique, SCSS et Bootstrap", "Interfaces statiques structurées"),
        ("Responsive", "Breakpoints, grille et adaptation des images", "Utilisation multi-écrans"),
        ("Accessibilité", "Alternatives textuelles et attributs ARIA", "Interface plus inclusive"),
        ("Collaboration", "Branche personnelle, commits et pull requests", "Travail collectif traçable"),
    ])

    doc.add_heading("1. Contexte du projet", level=1)
    doc.add_paragraph(
        "Dans le cadre de la formation DWWM, le projet consistait à concevoir puis à réaliser un site vitrine "
        "pour une activité de photographie professionnelle. Le site devait présenter l'univers du photographe, "
        "ses prestations, ses galeries et ses contenus éditoriaux, tout en facilitant la prise de contact et la "
        "réservation d'une séance."
    )
    doc.add_paragraph(
        "La conception a été menée en équipe. Nous avons étudié les contenus attendus, organisé les pages et "
        "construit les maquettes sur Figma. Le développement a ensuite été réalisé par le même groupe à partir "
        "d'un dépôt GitHub commun."
    )

    doc.add_heading("Objectifs fonctionnels", level=2)
    add_three_col_table(doc, ["Objectif", "Ce que cela apporte", "Exemple dans le projet"], [
        ("Présenter les prestations", "Compréhension rapide de l'offre", "Mariage, maternité, portrait, sport, animalier et événements"),
        ("Valoriser les photographies", "Hiérarchie visuelle sobre", "Grands formats d'image et espaces aérés"),
        ("Adapter l'interface", "Navigation fluide sur ordinateur et mobile", "Maquettes desktop et mobile, règles responsive"),
        ("Créer des appels à l'action", "Faciliter la réservation", "Boutons vers contact, galerie ou blog"),
    ])

    doc.add_heading("2. Conception en équipe", level=1)
    doc.add_paragraph(
        "Avec l'équipe, j'ai participé à la définition des grandes zones de contenu et du parcours utilisateur. "
        "Cette étape nous a permis de hiérarchiser l'information avant de travailler les détails graphiques. "
        "Nous avons ensuite décliné les écrans dans Figma en distinguant les formats desktop et mobile."
    )
    add_image_with_caption(doc, images["media/image2.png"], "Wireframes des principales pages et déclinaisons desktop/mobile réalisés sur Figma.", 5.6)

    doc.add_heading("Maquettage responsive de la page d'accueil", level=2)
    doc.add_paragraph(
        "Dans ce projet réalisé en équipe, j'ai pris en charge la conception de la page d'accueil sur Figma, "
        "dans ses versions desktop et mobile. J'ai organisé les sections : bannière principale, présentation du "
        "photographe, prestations, témoignages, tarifs, contact et pied de page. J'ai adapté les dimensions des "
        "images, les textes et les espacements pour garantir une navigation claire sur différentes tailles d'écran."
    )
    add_image_with_caption(doc, images["media/image4.png"], "Maquettes desktop et mobile de la page d'accueil réalisées sur Figma.", 3.6)

    doc.add_heading("Prototypage et interactions", level=2)
    doc.add_paragraph(
        "Après les maquettes, j'ai réalisé le prototypage interactif sur Figma. J'ai relié les boutons, les éléments "
        "de navigation et les appels à l'action aux écrans correspondants afin de simuler le parcours utilisateur "
        "avant le développement."
    )
    add_image_with_caption(doc, images["media/image5.png"], "Création des interactions et du parcours utilisateur dans le prototype Figma.", 5.4)

    doc.add_heading("Direction artistique", level=2)
    doc.add_paragraph(
        "L'identité visuelle repose sur une palette douce et chaleureuse, adaptée à l'univers de la photographie "
        "de famille et d'événement. Le blanc structure les espaces, les tons rosés servent d'accent et les teintes "
        "brunes apportent du contraste."
    )
    add_image_with_caption(doc, images["media/image6.png"], "Palette de couleurs utilisée pour l'identité visuelle du site.", 3.0)
    doc.add_paragraph(
        "Playfair Display est utilisée pour les titres principaux afin de donner une dimension éditoriale. "
        "Montserrat améliore la lisibilité des textes courants et des éléments d'interface. Une police manuscrite, "
        "Miama, est employée ponctuellement pour renforcer la personnalité de certains titres."
    )

    doc.add_heading("3. Passage de la maquette au site web", level=1)
    doc.add_paragraph(
        "Après la phase de conception, nous avons poursuivi le développement du site. Le dépôt commun a été "
        "structuré autour de pages HTML, de feuilles de style CSS générées depuis SCSS et d'un dossier d'images "
        "et de polices. Bootstrap a été utilisé pour accélérer la mise en page responsive et fournir des composants "
        "d'interface fiables."
    )
    add_three_col_table(doc, ["Outil", "Contribution personnelle", "Apport professionnel"], [
        ("HTML5", "Structure sémantique des pages et des contenus", "Lisibilité, accessibilité et référencement"),
        ("SCSS", "Styles modulaires et classes structurées", "Maintenance et réutilisation"),
        ("Bootstrap 5", "Grille responsive, utilitaires et composant collapse", "Adaptation rapide aux écrans"),
        ("Git/GitHub", "Branche faten-modifications et pull requests", "Traçabilité et intégration collective"),
        ("Figma", "Maquettes desktop et mobile", "Référence visuelle partagée"),
    ])

    doc.add_heading("4. Contribution personnelle : illustration de contact", level=1)
    doc.add_paragraph(
        "L'illustration de contact constitue un appel à l'action placé à un moment stratégique du parcours. Elle "
        "associe une image, un titre, un texte court et un bouton de réservation. L'objectif était d'offrir une "
        "transition claire vers la prise de contact sans rompre l'identité visuelle du site."
    )
    doc.add_heading("Structure HTML", level=2)
    add_code_block(doc, [
        '<div class="illustration-contact container-fluid p-0">',
        '  <div class="illustration-contact__content ...">',
        '    <h2>Prêt à capturer vos plus beaux instants ?</h2>',
        '    <a class="btn illustration-contact__button">Réserver une séance</a>',
        '  </div>',
        '</div>',
    ])
    doc.add_heading("Mise en forme responsive", level=2)
    doc.add_paragraph(
        "J'ai organisé les styles avec des classes de type BEM, comme illustration-contact__card et "
        "illustration-contact__button. La largeur de la carte est limitée pour conserver une bonne lisibilité, "
        "tandis que les dimensions et la typographie sont réduites sous 768 px. L'image passe en arrière-plan "
        "sur mobile afin de préserver la composition."
    )
    add_code_block(doc, [
        "@media (max-width: 767.98px) {",
        "  .illustration-contact {",
        "    background-image: url('../assets/img/contact__image.png');",
        "    background-size: cover;",
        "  }",
        "  .illustration-contact__card {",
        "    width: min(312px, calc(100% - 32px));",
        "  }",
        "}",
    ])
    add_label_para(doc, "Preuve GitHub : ", "création des fichiers HTML et SCSS le 5 juin 2026, ajout de l'image contact__image.png, puis corrections successives sur la branche faten-modifications.")

    doc.add_heading("5. Contribution personnelle : page Blog", level=1)
    doc.add_paragraph(
        "La page Blog devait présenter les contenus éditoriaux du photographe et permettre à l'utilisateur de "
        "retrouver rapidement un article. J'ai développé le hero, la navigation par catégories, les cartes "
        "d'articles et la barre latérale."
    )
    add_three_col_table(doc, ["Bloc", "Réalisation", "Objectif"], [
        ("Hero", "Image, voile de contraste, titre et introduction", "Entrée visuelle claire"),
        ("Catégories", "Navigation desktop et filtres repliables sur mobile", "Parcours adapté au terminal"),
        ("Articles", "Image, catégorie, titre, résumé, date et lien", "Lecture rapide des contenus"),
        ("Sidebar", "Recherche, catégories et articles populaires", "Accès direct aux contenus utiles"),
    ])
    doc.add_heading("Responsive et hiérarchie", level=2)
    doc.add_paragraph(
        "Sur ordinateur, la liste d'articles occupe neuf colonnes et la barre latérale trois colonnes. Sur les "
        "écrans plus étroits, les blocs passent en pleine largeur. La navigation desktop est masquée sous 768 px "
        "et remplacée par un bouton de filtres utilisant le composant collapse de Bootstrap."
    )
    add_code_block(doc, [
        '<section class="col-12 col-lg-9">...</section>',
        '<aside class="col-12 col-lg-3">...</aside>',
        '<button data-bs-toggle="collapse"',
        '        data-bs-target="#mobileFilters"',
        '        aria-controls="mobileFilters">Afficher les filtres</button>',
    ])
    add_label_para(doc, "Preuve GitHub : ", "création de la page Blog le 8 juin 2026, enrichissement progressif de son contenu et de ses styles, puis intégration par pull requests.")

    doc.add_heading("6. Qualité de l'interface", level=1)
    add_label_para(doc, "Responsive design : ", "les règles SCSS utilisent plusieurs points de rupture : 991,98 px pour les tablettes, 767,98 px pour les mobiles et 399,98 px pour les très petits écrans. Les tailles de titre, les espacements, la position des images et les menus sont ajustés à chaque étape.")
    add_label_para(doc, "Accessibilité : ", "les images de contenu disposent d'un attribut alt descriptif. Les zones de navigation sont nommées avec aria-label. Le bouton de filtres renseigne aria-expanded et aria-controls, et les icônes décoratives sont masquées avec aria-hidden.")
    add_label_para(doc, "Maintenabilité : ", "la convention de nommage des classes permet de repérer rapidement le bloc, l'élément et son rôle. Les styles sont regroupés par composant et les valeurs responsive sont centralisées dans les media queries.")

    doc.add_heading("Tests réalisés dans le cadre du travail", level=2)
    add_three_col_table(doc, ["Contrôle", "Méthode", "Résultat attendu"], [
        ("Conformité visuelle", "Comparaison avec les maquettes desktop et mobile", "Interface fidèle à la maquette"),
        ("Responsive", "Redimensionnement de la fenêtre", "Aucun débordement horizontal"),
        ("Interactions", "Vérification des filtres mobiles, liens et images", "Parcours utilisable"),
        ("Focus", "Contrôle des états de focus et de survol", "Navigation plus accessible"),
    ])

    doc.add_heading("7. Collaboration et traçabilité", level=1)
    doc.add_paragraph(
        "Le projet a été suivi dans le dépôt Juanrojasdc/projet-site-photo. J'ai travaillé sur la branche "
        "faten-modifications afin d'isoler mes changements et d'éviter de modifier directement la branche principale. "
        "Cette méthode permettait à l'équipe de relire et d'intégrer les contributions de manière contrôlée."
    )
    add_three_col_table(doc, ["Date", "Contribution personnelle", "Apport"], [
        ("5 juin", "Création de l'illustration Contact en HTML et SCSS", "Premier composant personnel"),
        ("5 juin", "Ajout de contact__image.png et corrections responsive", "Amélioration visuelle"),
        ("8 juin", "Création de la page Le Blog et de ses styles", "Nouvelle page complète"),
        ("8-9 juin", "Corrections HTML/SCSS et enrichissement de la page", "Itérations et stabilisation"),
        ("PR #2 à #5", "Propositions d'intégration depuis faten-modifications", "Collaboration traçable"),
    ])
    doc.add_paragraph("Dépôt de référence : github.com/Juanrojasdc/projet-site-photo")

    doc.add_heading("Difficultés rencontrées et solutions", level=2)
    doc.add_paragraph(
        "L'adaptation des visuels aux différentes tailles d'écran demandait plusieurs ajustements d'object-fit, "
        "d'object-position et de dimensions. J'ai résolu ce point en testant progressivement les points de rupture "
        "et en définissant des règles spécifiques pour les petits écrans."
    )
    doc.add_paragraph(
        "Le travail collectif nécessitait également de maintenir ma branche à jour et de proposer des modifications "
        "limitées à mes fichiers. Les pull requests et les commits successifs m'ont permis de conserver une trace "
        "claire des corrections."
    )

    doc.add_heading("8. Réponses synthétiques pour validation", level=1)
    doc.add_heading("Tâches réalisées et conditions", level=2)
    doc.add_paragraph(
        "J'ai participé avec mon équipe à la conception des maquettes responsive d'un site de photographie dans Figma. "
        "Ensemble, nous avons défini l'organisation des contenus, la direction artistique, les versions desktop et "
        "mobile et les principales pages. J'ai ensuite participé au développement collectif du site LA Photographie33. "
        "Ma contribution identifiable a porté sur l'illustration de contact et sur la page Blog, développées en HTML5, "
        "SCSS et Bootstrap, puis intégrées au moyen de Git et GitHub."
    )
    doc.add_heading("Moyens utilisés", level=2)
    doc.add_paragraph(
        "Figma pour la conception des écrans ; HTML5 pour la structure ; SCSS et CSS3 pour les styles ; Bootstrap 5 "
        "et Bootstrap Icons pour la grille, les utilitaires et les composants responsive ; Git et GitHub pour la branche "
        "personnelle, les commits et les pull requests ; images et polices intégrées au projet pour respecter l'identité visuelle."
    )
    doc.add_heading("Avec qui avez-vous travaillé ?", level=2)
    doc.add_paragraph(
        "J'ai participé, avec mon équipe de trois personnes, à la conception des maquettes responsive du site "
        "LA Photographie33 sur Figma. Lors de la phase de conception, j'ai réalisé les maquettes des pages d'accueil "
        "et de contact dans leurs versions desktop et mobile. J'ai ensuite participé au développement collectif du site."
    )
    doc.add_heading("Contexte", level=2)
    add_info_table(doc, [
        ("Organisme", "AFPA Bègles"),
        ("Période d'exercice", "Du 23/03/2026 au 18/12/2026"),
        ("Cadre", "Projet pédagogique réalisé pendant la formation Développeur Web et Web Mobile"),
    ])

    doc.add_heading("Axes d'amélioration", level=2)
    add_three_col_table(doc, ["Axe", "Action prévue", "Intérêt"], [
        ("Validation", "Renforcer les tests automatisés et les contrôles HTML/CSS", "Fiabiliser l'intégration"),
        ("Performance", "Optimiser davantage les images et mesurer les performances", "Améliorer le chargement"),
        ("Documentation", "Documenter les composants partagés et centraliser les variables SCSS", "Faciliter la maintenance"),
        ("Accessibilité", "Approfondir les tests clavier et lecteur d'écran", "Rendre l'interface plus inclusive"),
    ])

    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "Cette réalisation confirme ma capacité à concevoir une interface, à développer des pages responsive et à "
        "collaborer dans un projet versionné. Le travail réalisé sur LA Photographie33 montre une contribution "
        "identifiable, vérifiable et reliée aux compétences front-end attendues dans le titre DWWM."
    )

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
